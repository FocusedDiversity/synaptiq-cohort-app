"""
Convert cohort-app-overview.md to a formatted Word document.
Run: python docs/build_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches

def rgb(hex_str: str) -> RGBColor:
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE  = Path(__file__).parent / "cohort-app-overview.md"
OUT_FILE = Path(__file__).parent / "cohort-app-overview.docx"

# ── colour palette ──────────────────────────────────────────────────────────
BRAND_BLUE   = "005FAF"
BRAND_TEAL   = "008C8C"
LIGHT_GREY   = "F2F2F2"
MID_GREY     = "D9D9D9"
DARK_GREY    = "404040"
WHITE        = "FFFFFF"


def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def set_cell_border(cell, border_colour="AAAAAA"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), border_colour)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            run.font.color.rgb = rgb(WHITE)
            p.paragraph_format.space_before = Pt(18)
        elif level == 2:
            run.font.color.rgb = rgb(BRAND_BLUE)
            p.paragraph_format.space_before = Pt(14)
        elif level == 3:
            run.font.color.rgb = rgb(BRAND_TEAL)
            p.paragraph_format.space_before = Pt(10)
        elif level == 4:
            run.font.color.rgb = rgb(DARK_GREY)
    return p


def shade_h1(doc, text):
    """H1 with blue background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "005FAF")
    pPr.append(shd)
    run = p.add_run(text)
    run.bold            = True
    run.font.size       = Pt(18)
    run.font.color.rgb  = rgb(WHITE)
    return p


def add_table(doc, headers, rows):
    col_count = len(headers)
    table     = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, BRAND_BLUE)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold           = True
        run.font.color.rgb = rgb(WHITE)
        run.font.size      = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg    = LIGHT_GREY if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row):
            cell = cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()


def add_code_block(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    # bold the **...** parts
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold      = True
            run.font.size = Pt(10)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)


# ── parse markdown ──────────────────────────────────────────────────────────

def parse_table(lines):
    """Parse a markdown table and return (headers, rows)."""
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows    = []
    for line in lines[2:]:  # skip separator
        if not line.strip() or set(line.strip().replace("|", "").replace("-", "").replace(" ", "")) == set():
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        # pad/trim to header count
        while len(cells) < len(headers):
            cells.append("")
        rows.append(cells[:len(headers)])
    return headers, rows


def build_doc(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    lines      = md_path.read_text(encoding="utf-8").splitlines()
    i          = 0
    code_buf   = []
    in_code    = False
    table_buf  = []
    in_table   = False

    while i < len(lines):
        line = lines[i]

        # ── code block ──────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, "\n".join(code_buf))
                in_code  = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── markdown table ──────────────────────────────────────────────────
        if line.startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        else:
            if table_buf:
                headers, rows = parse_table(table_buf)
                add_table(doc, headers, rows)
                table_buf = []

        # ── headings ────────────────────────────────────────────────────────
        if line.startswith("# SECTION"):
            shade_h1(doc, line.lstrip("# ").strip())
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # strip italic markers from h3
            text  = re.sub(r"\*(.+?)\*", r"\1", text)
            add_heading(doc, text, level)
            i += 1
            continue

        # ── horizontal rule ─────────────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # ── bullet ──────────────────────────────────────────────────────────
        bm = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bm:
            indent = len(bm.group(1)) // 2
            add_bullet(doc, bm.group(2), level=indent)
            i += 1
            continue

        # ── blank line ───────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────────────────
        # strip inline bold/italic for display
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        clean = re.sub(r"\*(.+?)\*",     r"\1", clean)
        clean = re.sub(r"`(.+?)`",        r"\1", clean)
        p = doc.add_paragraph()
        # re-apply bold spans
        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold      = True
                run.font.size = Pt(10)
            else:
                part = re.sub(r"\*(.+?)\*", r"\1", part)
                part = re.sub(r"`(.+?)`",    r"\1", part)
                run  = p.add_run(part)
                run.font.size = Pt(10)
        i += 1

    # flush any trailing table
    if table_buf:
        headers, rows = parse_table(table_buf)
        add_table(doc, headers, rows)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc(MD_FILE, OUT_FILE)
